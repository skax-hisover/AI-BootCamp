"""Knowledge document loading and chunk preparation."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx"}
REQUIRED_SIDECAR_META_FIELDS = ("collected_at", "source_url", "curator", "license")


def _load_sidecar_metadata(path: Path) -> dict[str, str]:
    """Load optional *.meta.json sidecar and keep only required metadata fields."""
    sidecar = path.with_name(f"{path.name}.meta.json")
    if not sidecar.exists():
        print(f"[WARN] Missing metadata sidecar for {path.name}: {sidecar.name}")
        return {}
    try:
        payload = json.loads(sidecar.read_text(encoding="utf-8"))
    except Exception as exc:
        print(f"[WARN] Invalid metadata sidecar for {path.name}: {exc}")
        return {}
    if not isinstance(payload, dict):
        print(f"[WARN] Invalid metadata sidecar for {path.name}: root must be object")
        return {}

    merged: dict[str, str] = {}
    missing_fields: list[str] = []
    for field in REQUIRED_SIDECAR_META_FIELDS:
        value = payload.get(field)
        if isinstance(value, str) and value.strip():
            merged[field] = value.strip()
        else:
            missing_fields.append(field)
    if missing_fields:
        missing = ", ".join(missing_fields)
        print(f"[WARN] Incomplete metadata sidecar for {path.name}: missing [{missing}]")
    return merged


def _infer_root_category_from_filename(path: Path) -> str:
    """Infer category for root-level files to avoid uncategorized overuse."""
    stem = path.stem.lower()
    filename = path.name.lower()
    text = f"{stem} {filename}"

    if any(token in text for token in ("job_posting", "posting", "공고", "채용", "job_market")):
        return "job_postings"
    if any(token in text for token in ("jd", "job_description", "직무기술", "job_desc")):
        return "jd"
    if any(token in text for token in ("interview", "면접")):
        return "interview_guides"
    if any(token in text for token in ("portfolio", "포트폴리오", "resume", "이력서")):
        return "portfolio_examples"
    return "uncategorized"


def _read_text_file(path: Path) -> list[dict[str, object]]:
    text = path.read_text(encoding="utf-8")
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    if not blocks:
        return []
    return [
        {"content": block, "metadata": {"section_type": "paragraph", "section_index": idx + 1}}
        for idx, block in enumerate(blocks)
    ]


def _read_csv_file(path: Path) -> list[dict[str, object]]:
    df = pd.read_csv(path).fillna("")
    rows = []
    for idx, row in enumerate(df.values.tolist(), start=1):
        row_text = " | ".join(map(str, row)).strip()
        if row_text:
            rows.append(
                {
                    "content": row_text,
                    "metadata": {"section_type": "csv_row", "row_number": idx},
                }
            )
    return rows


def _read_pdf_file(path: Path) -> list[dict[str, object]]:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing dependency for PDF loader: install pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[dict[str, object]] = []
    for page_idx, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages.append(
                {
                    "content": page_text,
                    "metadata": {"section_type": "pdf_page", "page_number": page_idx},
                }
            )
    return pages


def _read_docx_file(path: Path) -> list[dict[str, object]]:
    try:
        from docx import Document as DocxDocument
    except ModuleNotFoundError as exc:
        raise RuntimeError("Missing dependency for DOCX loader: install python-docx") from exc

    doc = DocxDocument(str(path))
    paragraphs: list[dict[str, object]] = []
    for para_idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            paragraphs.append(
                {
                    "content": text,
                    "metadata": {"section_type": "docx_paragraph", "paragraph_number": para_idx},
                }
            )
    return paragraphs


def _read_xlsx_file(path: Path) -> list[dict[str, object]]:
    xls = pd.ExcelFile(path)
    blocks: list[dict[str, object]] = []
    for sheet in xls.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet).fillna("")
        for row_idx, row in enumerate(df.values.tolist(), start=1):
            row_text = " | ".join(map(str, row)).strip()
            if row_text:
                blocks.append(
                    {
                        "content": row_text,
                        "metadata": {
                            "section_type": "xlsx_row",
                            "sheet_name": sheet,
                            "row_number": row_idx,
                        },
                    }
                )
    return blocks


def iter_source_files(knowledge_dir: Path) -> Iterable[Path]:
    for path in sorted(knowledge_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path


def load_documents_with_report(knowledge_dir: Path) -> tuple[list[Document], list[dict[str, str]]]:
    docs: list[Document] = []
    failures: list[dict[str, str]] = []
    for path in iter_source_files(knowledge_dir):
        try:
            suffix = path.suffix.lower()
            if suffix == ".csv":
                sections = _read_csv_file(path)
            elif suffix == ".pdf":
                sections = _read_pdf_file(path)
            elif suffix == ".docx":
                sections = _read_docx_file(path)
            elif suffix == ".xlsx":
                sections = _read_xlsx_file(path)
            else:
                sections = _read_text_file(path)
        except Exception as exc:
            # Keep service available even if one file is malformed or parser dependency is missing.
            print(f"[WARN] Failed to load {path.name}: {exc}")
            relative = path.relative_to(knowledge_dir).as_posix()
            failures.append({"file": relative, "error": str(exc)})
            continue

        if not sections:
            continue

        relative = path.relative_to(knowledge_dir)
        category = (
            relative.parts[0]
            if len(relative.parts) > 1
            else _infer_root_category_from_filename(path)
        )
        sidecar_meta = _load_sidecar_metadata(path)
        for section in sections:
            content = str(section.get("content", "")).strip()
            if not content:
                continue
            section_meta = section.get("metadata", {})
            metadata = {
                "source": str(path),
                "filename": path.name,
                "relative_path": str(relative).replace("\\", "/"),
                "category": category,
            }
            if isinstance(section_meta, dict):
                metadata.update(section_meta)
            metadata.update(sidecar_meta)

            docs.append(
                Document(
                    page_content=content,
                    metadata=metadata,
                )
            )
    return docs, failures


def load_documents(knowledge_dir: Path) -> list[Document]:
    docs, _ = load_documents_with_report(knowledge_dir)
    return docs


def chunk_documents(
    docs: list[Document],
    chunk_size: int,
    chunk_overlap: int,
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " "],
    )
    chunks = splitter.split_documents(docs)
    per_source_counter: dict[str, int] = {}
    for idx, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = idx
        source = str(chunk.metadata.get("source", "unknown"))
        per_source_counter[source] = per_source_counter.get(source, 0) + 1
        chunk.metadata["chunk_index_in_source"] = per_source_counter[source]
    return chunks
